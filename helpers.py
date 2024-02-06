import cohere
import pandas as pd
from youtube_transcript_api import YouTubeTranscriptApi
from googleapiclient.discovery import build
import time
import docx
from docx import Document
import io

# Initialize YouTube API
def youtube_service(api_key):
    return build('youtube', 'v3', developerKey=api_key)

# Search for YouTube videos
def search_youtube(youtube, query):
    request = youtube.search().list(part="snippet", maxResults=5, q=query, type="video")
    response = request.execute()
    return response.get('items', [])


# Function to fetch and process YouTube transcripts, languages=['en']
def fetch_youtube_transcript(video_id):
    try:
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
    except:
        try:
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en-US'])
        except:
            try:
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en-GB'])
            except:
                return "No Transcript Found"
    if isinstance(transcript_list, str):
        return "No Transcript Found"
    df = pd.DataFrame(transcript_list)
    return df

# Function to summarize a batch of text segments
def summarize_batch(cohere_client, combined_text, summary_length="short"):
    return cohere_client.summarize(text=combined_text, model='command', length=summary_length)

# Function to process transcript data in batches for summarization
def process_transcript_for_summaries(cohere_client, transcript_df, batch_size=50):

    for start in range(0, len(transcript_df), batch_size):
        run_start = time.time()
        end = start + batch_size
        batch_segments = transcript_df['text'][start:end].tolist()

        try:
            combined_text = ' '.join(batch_segments)
            summary_text = summarize_batch(cohere_client, combined_text)
            start_time = transcript_df['start'].iloc[start]
            summary = {'start': start_time, 'summary': summary_text.summary, 'original':combined_text, 'run_time':time.time()-run_start}

            yield summary

        except Exception as e:
            print('Error', e)


# Main execution function
def summarize(conn, video_id):
    # Initialize Cohere client
    cohere_client = conn

    # Fetch and process YouTube transcript
    transcript_df = fetch_youtube_transcript(video_id)

    if isinstance(transcript_df, str):
        return "No trancript Found"
    # Process the transcript for summaries
    summary_df = process_transcript_for_summaries(cohere_client, transcript_df)

    # Return or display the summary DataFrame
    return {'Transcript':transcript_df,
            'Summary':summary_df}

def create_word_document(video_details, summaries):
    # Create a new Document
    doc = Document()
    
    # Set the title of the document
    doc.add_heading(video_details['video_title'], level=0)

    doc.add_paragraph(f"Video Link : https://www.youtube.com/watch?v={video_details['video_id']}")
    
    # Iterate through the summaries dictionary
    # Assuming 'summaries' is a dictionary with start times as keys and summary text as values
    for start_time, summary in summaries.items():
        if start_time == '-1':
            continue

        # Add a heading for each start time
        doc.add_heading(f'Start Time: {round(start_time)} s', level=1)
        # Add the summary text as a new paragraph
        doc.add_paragraph(summary)
    
     # Save the document to a BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)  # Go back to the beginning of the BytesIO object after saving
    
    return doc_io

def conversation_query(conn, conversation, query):
    """
    The chat endpoint allows users to have conversations with a Large Language Model (LLM) from Cohere. 
    Users can send messages as part of a persisted conversation using the conversation_id parameter, 
    or they can pass in their own conversation history using the chat_history parameter.
    The endpoint features additional parameters such as connectors and documents that enable conversations
    enriched by external knowledge. We call this "Retrieval Augmented Generation", or "RAG".
    """
    response = conn.chat(
    chat_history= conversation,
    message= query,
    # perform web search before answering the question. You can also use your own custom connector.
    connectors=[{"id": "web-search"}] 
    )
    return response.text

def generate(conn, prompt, model, max_tokens):
    """
    This endpoint generates realistic text conditioned on a given input.
    """
    return conn.generate(
                    model=model,
                    prompt = prompt,
                    max_tokens = max_tokens
                    )
def detect_language(conn, texts):
    """
    This endpoint identifies which language each of the provided texts is written in.
    """
    return conn.detect_language(
                                texts=texts
                                ).results

def relevance_ranked(conn, docs, query):
    """
    This endpoint takes in a query and a list of texts and produces an ordered array with each text assigned a relevance score.
    """

    return conn.rerank(
            model = 'rerank-english-v2.0',
            query = query,
            documents = docs,
            top_n = 3,
            )

def tokenize(conn, text):
    """
    This endpoint splits input text into smaller units called tokens using byte-pair encoding (BPE).
    To learn more about tokenization and byte pair encoding, see the tokens page.
    """
    return conn.tokenize(
            text=text,
            model='command' 
            )

def detokenize(conn, tokens):
    """
    This endpoint takes tokens using byte-pair encoding and returns their text representation.
    To learn more about tokenization and byte pair encoding, see the tokens page.
    """
    return conn.detokenize(
            tokens=tokens,
            model="command"
            )




